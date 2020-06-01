#!/usr/bin/env python3

#pip install docx
#pip install pandas
#pip install xlrd

from docx import Document
from docx.shared import Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT

import pandas as pd
import math
import argparse
import sys




# input excel file, return dataframe or error
def obtainExcelDf(filepath):
    try:
        # read excel file
        df=pd.read_excel(filepath,header=None)
    except:
        return "File does not exist"
    
    
    if(df.shape[0]==0 or df.shape[1]==0):
        return "Not enough data in excel sheet"
    
    # drop rows with any NaN
    df.dropna(inplace=True)
    df.reset_index(drop=True,inplace=True)
    
    # set column names
    df.columns=df.iloc[0];
    df.drop(0,inplace=True)
    
    # sort by lastname (last column)
    df.sort_values(by=df.columns[-1],inplace=True)
    
    return df

#########################################################

# merge columns in a dataframe. display result as a list
def mergeDfColsAndList(df,listColNames):
    listSeries=[];
    try:
        for colName in listColNames:
            listSeries.append(df[colName])
    except:
        return "Invalid column names. Can't merge columns."
    
    if(len(listSeries)>0):
        returnList=listSeries[0].astype('str')
        for i in range(1,len(listSeries)):
            returnList=returnList+", "+listSeries[i].astype('str')
        return returnList.to_list()


#########################################################
    
      

# helper functions (assume input values are valid)

# delete n rows from table in document
def deleteLastNRows(table,n):
    
    tbl = table._tbl
    for i in range (n):
        row=table.rows[-1]
        tr = row._tr
        tbl.remove(tr)
        
# add n rows to document table and format
def addNRows(table,n):
    lastRow=table.rows[-1]
    
    for i in range (n):
        row=table.add_row()
        row.height=lastRow.height
        
        cells=row.cells
        for cell in cells:
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER



#########################################################            
                      


# create word document of labels
def createLabelDoc(contentList, template, fileName="labels.docx", columns=3, font='Arial', fontSize=22):
    
    # invalid types
    if(type(contentList) is not list or type(template)is not str or type(fileName) is not str or type(columns) is not int
       or type(font) is not str or type(fontSize)is not int):
        return "Invalid argument type"

    
    # copy template into new doc
    try:
        templateDocument = Document(template)
    except:
        return "Invalid template file. Please make sure it is a .docx file."
    document=Document()
    for elem in templateDocument.element.body:
        document.element.body.append(elem)

    
    # if no table is in template
    if(len(document.tables)==0):
        return "Table not found in template file"
          
        
    table=document.tables[0]
    
    
    # total label needed
    labelsNeeded=len(contentList)
    # rows needed in table
    rowsNeeded=math.ceil(len(contentList)/columns)
    
    
    # handles for rows like   | e | e | e | e |   or   | e || e || e || e |   or   | e ||||| e ||||| e ||||| e |
    # checks for matching table column count. # of cells inbetween column entry is multFactor-1    
    multFactor=1    
    if((len(table.columns)-columns)%(columns-1)==0):
        multFactor+=int((len(table.columns)-columns)/(columns-1))
    else:
        return"Columns in template does not match for column value: {}".format(columns)
    
    
    
    # add/subtract rows to obtain rows needed
    difference=rowsNeeded-len(table.rows)
    if(difference>0):
        addNRows(table,difference)
    elif(difference<0):
        deleteLastNRows(table,0-difference)

    
    
    # add entries to table
    for row in range(rowsNeeded):
        for col in range(columns): 

            # current contentList index
            index=row*columns+col;

            # if index is within the length of contentList, append name to table
            if(index<labelsNeeded):
                name=contentList[index]
                
                cell=table.cell(row,col*multFactor)
                
                p=cell.paragraphs[0]
                p.text= name
                p.runs[0].font.name=font
                p.runs[0].font.size=Pt(fontSize)
                
    document.save(fileName)


    
#########################################################

# excel to df, merge columns of df, create document with labels
def createNameLabelDoc(excelFile,template,fileName="labels.docx",columns=3,font='Arial',fontSize=22, *args): 
    
    df=obtainExcelDf(excelFile)
    if(type(df) is str):
        return df

    nameList=mergeDfColsAndList(df,*args)
    if(type(nameList) is str):
        return nameList

    docError=createLabelDoc(nameList,template,fileName,columns,font,fontSize)
    if(type(docError) is str):
        return docError


#########################################################

argumentIndex={
    'excelFile':0,
    'template':1,
    'fileName':2,
    'columns':3,
    'font':4,
    'fontSize':5,
    'args':6
}




def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('-e', '--excelFile',nargs=1,required=True, help='path to the excel file with the data')
    parser.add_argument('-t', '--template',nargs=1,required=True, help='path to the label word template')
    parser.add_argument('-f', '--fileName',nargs='?',default='labels.docx',help='output file name')
    parser.add_argument('-c', '--columns', type=int,nargs='?',default=3, help='number of columns to fill in template')
    parser.add_argument('-o', '--font',nargs='?',default='Arial', help='font to fill template')
    parser.add_argument('-s', '--fontSize', type=int,nargs='?',default=22, help='font size to fill template')
    parser.add_argument('-a','--args', nargs='*',help='list of column names to merge in excel file')
    
    dictArg={}
    for key, value in parser.parse_args()._get_kwargs():
        if(type(value) is list and argumentIndex[key]!=6):
            dictArg[argumentIndex[key]]=value[0]
        else:
            dictArg[argumentIndex[key]]=value
    return dictArg
 

if __name__ == '__main__':
    dictArgs=main()
    error=createNameLabelDoc(dictArgs[0],dictArgs[1],dictArgs[2],dictArgs[3],dictArgs[4],dictArgs[5], dictArgs[6])
    if(error!=None):
        print(error)




